#!/usr/bin/env python3
"""
Word 转 PDF 插件

将 Word 文档转换为 PDF 文件，优先使用 docx2pdf 库（需要系统安装 Microsoft Word 或 LibreOffice），
如果不可用则降级到 python-docx + reportlab 纯 Python 实现。
"""

from pathlib import Path
from cedar.utils import print


def execute(word_path: str, output_path: str = None) -> tuple:
    """
    将 Word 文档转换为 PDF 文件
    
    Args:
        word_path: Word 文件路径
        output_path: 输出 PDF 文件路径，如果为 None 则自动生成
        
    Returns:
        tuple: (success: bool, message: str, output_path: str)
    """
    print(f"[DEBUG] [WordToPdfPlugin] 开始转换 Word 到 PDF: {word_path}")
    word_path = Path(word_path)
    if not word_path.exists():
        print(f"[ERROR] [WordToPdfPlugin] Word 文件不存在: {word_path}")
        return False, "Word 文件不存在", None
    
    # 验证文件格式
    if not word_path.suffix.lower() in ['.docx', '.doc']:
        print(f"[ERROR] [WordToPdfPlugin] 文件格式不支持: {word_path.suffix}")
        return False, f"不支持的文件格式: {word_path.suffix}，仅支持 .docx 和 .doc", None
    
    # 如果没有指定输出路径，自动生成
    if output_path is None:
        output_path = word_path.parent / f"{word_path.stem}.pdf"
    else:
        output_path = Path(output_path)
    
    print(f"[DEBUG] [WordToPdfPlugin] 输出路径: {output_path}")
    
    # 如果文件已存在，添加序号
    counter = 1
    original_path = output_path
    while output_path.exists():
        output_path = original_path.parent / f"{original_path.stem}_{counter}{original_path.suffix}"
        counter += 1
        print(f"[DEBUG] [WordToPdfPlugin] 文件已存在，尝试新名称: {output_path}")
    
    # 优先尝试使用 docx2pdf 库进行转换（需要系统安装 Microsoft Word 或 LibreOffice）
    try:
        from docx2pdf import convert
        print(f"[DEBUG] [WordToPdfPlugin] docx2pdf 库已导入，开始转换...")
        print(f"[DEBUG] [WordToPdfPlugin] 使用 docx2pdf 转换（需要系统 Office 软件）...")
        
        # docx2pdf 需要绝对路径
        convert(str(word_path.absolute()), str(output_path.absolute()))
        print(f"[DEBUG] [WordToPdfPlugin] docx2pdf 转换完成")
        
        # 验证输出文件是否生成
        if not output_path.exists():
            print(f"[ERROR] [WordToPdfPlugin] PDF 文件未生成: {output_path}")
            raise Exception("输出文件未生成")
        
        file_size = output_path.stat().st_size
        print(f"[DEBUG] [WordToPdfPlugin] ✓ PDF 文件已生成: {output_path}, 大小={file_size} 字节")
        return True, f"转换成功：{output_path.name}", str(output_path)
        
    except ImportError:
        print(f"[WARN] [WordToPdfPlugin] docx2pdf 库未安装，尝试降级方案...")
        # 降级到 python-docx + reportlab
        return _convert_with_python_docx(word_path, output_path)
    except Exception as e:
        print(f"[WARN] [WordToPdfPlugin] docx2pdf 转换失败: {e}，尝试降级方案...")
        # 降级到 python-docx + reportlab
        return _convert_with_python_docx(word_path, output_path)


def _convert_with_python_docx(word_path: Path, output_path: Path) -> tuple:
    """
    使用 python-docx + reportlab 进行转换（降级方案）
    
    注意：此方案可能无法完全保留原始格式，特别是表格、复杂布局等
    
    Args:
        word_path: Word 文件路径
        output_path: 输出 PDF 文件路径
        
    Returns:
        tuple: (success: bool, message: str, output_path: str)
    """
    print(f"[DEBUG] [WordToPdfPlugin] 开始使用 python-docx + reportlab 转换...")
    
    try:
        # 尝试导入 python-docx
        try:
            from docx import Document
            print(f"[DEBUG] [WordToPdfPlugin] python-docx 库已导入")
        except ImportError:
            print(f"[ERROR] [WordToPdfPlugin] python-docx 库未安装")
            return False, "转换失败：请安装 python-docx 库（pip install python-docx）", None
        
        # 尝试导入 reportlab
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
            print(f"[DEBUG] [WordToPdfPlugin] reportlab 库已导入")
        except ImportError:
            print(f"[ERROR] [WordToPdfPlugin] reportlab 库未安装")
            return False, "转换失败：请安装 reportlab 库（pip install reportlab），或运行 pip install -r requirements.txt 安装所有依赖", None
        
        print(f"[DEBUG] [WordToPdfPlugin] 依赖库已导入，开始读取 Word 文档...")
        
        # 读取 Word 文档
        doc = Document(str(word_path))
        print(f"[DEBUG] [WordToPdfPlugin] Word 文档已读取，段落数={len(doc.paragraphs)}")
        
        # 创建 PDF 文档
        pdf_doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # 获取样式
        styles = getSampleStyleSheet()
        
        # 创建自定义样式
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=12,
            alignment=TA_LEFT
        )
        
        heading1_style = ParagraphStyle(
            'CustomHeading1',
            parent=styles['Heading1'],
            fontSize=14,
            spaceAfter=10,
            alignment=TA_LEFT
        )
        
        heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontSize=12,
            spaceAfter=8,
            alignment=TA_LEFT
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=6,
            alignment=TA_LEFT
        )
        
        print(f"[DEBUG] [WordToPdfPlugin] PDF 样式已创建，开始提取内容...")
        
        # 提取内容并构建 PDF 元素
        story = []
        paragraph_count = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            paragraph_count += 1
            print(f"[DEBUG] [WordToPdfPlugin] 处理段落 {paragraph_count}: {text[:50]}...")
            
            # 根据样式选择段落样式
            style_name = para.style.name if para.style else 'Normal'
            
            if 'Title' in style_name:
                story.append(Paragraph(text, title_style))
            elif 'Heading 1' in style_name or '标题 1' in style_name:
                story.append(Paragraph(text, heading1_style))
            elif 'Heading 2' in style_name or '标题 2' in style_name:
                story.append(Paragraph(text, heading2_style))
            else:
                story.append(Paragraph(text, normal_style))
            
            story.append(Spacer(1, 0.2 * inch))
        
        print(f"[DEBUG] [WordToPdfPlugin] 内容提取完成，共 {paragraph_count} 个段落，开始生成 PDF...")
        
        # 生成 PDF
        pdf_doc.build(story)
        print(f"[DEBUG] [WordToPdfPlugin] PDF 生成完成")
        
        # 验证输出文件
        if not output_path.exists():
            print(f"[ERROR] [WordToPdfPlugin] PDF 文件未生成: {output_path}")
            return False, "转换失败：输出文件未生成", None
        
        file_size = output_path.stat().st_size
        print(f"[DEBUG] [WordToPdfPlugin] ✓ PDF 文件已生成: {output_path}, 大小={file_size} 字节")
        return True, f"转换成功（降级方案）：{output_path.name}", str(output_path)
        
    except Exception as e:
        print(f"[ERROR] [WordToPdfPlugin] 降级方案转换失败: {e}")
        return False, f"转换失败：{str(e)}", None


if __name__ == "__main__":
    """测试代码"""
    print("=" * 60)
    print("[TEST] 测试 Word 转 PDF 插件")
    print("=" * 60)
    
    # 注意：需要提供一个真实的 Word 文件路径进行测试
    test_word = Path("/Users/zhangsong/Desktop/code/cedar_dev/mac-commondX/test/pdf/test_mermaid.docx")
    
    if test_word.exists():
        success, msg, output_path = execute(str(test_word))
        print(f"[TEST] 转换结果: {msg}")
        
        if success:
            print(f"[TEST] ✓ 测试通过，输出文件: {output_path}")
        else:
            print(f"[TEST] ✗ 转换失败")
    else:
        print(f"[TEST] 测试文件不存在: {test_word}")
